import io
from datetime import datetime
from io import BytesIO

import matplotlib.pyplot as plt
from django.utils.translation import gettext as _
from docx import Document

from energy.logic.aggregated_consumption.models import AggregationIntervalSeconds
from energy.logic.aggregated_consumption.parameters import (
    CommonQueryParameters,
    OneHourAggregationIntervalQueryParameters,
)
from energy.logic.aggregated_consumption.types import (
    ConsumptionRawAndFormatted, ConsumptionRawAndFormattedWithForecastRawAndFormatted,
    TotalConsumption,
)
from energy.logic.consumption_report.verbose_constants import \
    (
    get_aggregation_interval_verbose, get_hour_filtering_method_verbose,
)


class ReportCreator:
    # interval and corresponding consumption
    __CONSUMPTION_TABLE_COLUMNS_COUNT = 2
    # heading and total consumption
    __CONSUMPTION_TABLE_ADDITIONAL_ROWS_COUNT = 2

    def __init__(
            self, energy_consumption_raw_and_formatted_with_optional_forecast:
            ConsumptionRawAndFormattedWithForecastRawAndFormatted | ConsumptionRawAndFormatted,
            total_consumption: TotalConsumption,
            query_parameters: CommonQueryParameters
    ):
        self.__energy_consumption_raw_and_formatted_with_optional_forecast = \
            energy_consumption_raw_and_formatted_with_optional_forecast
        self.__query_parameters = query_parameters
        self.__total_consumption = total_consumption
        self.__report = Document()

    def create_report(self) -> BytesIO:
        self.__add_current_date_and_time()
        self.__add_report_parameters_info()
        self.__add_consumption_table()
        self.__add_chart()
        self.__add_project_manager()
        return self.__export_report_to_binary()

    def __export_report_to_binary(self) -> BytesIO:
        binary_report = BytesIO()
        self.__report.save(binary_report)
        binary_report.seek(io.SEEK_SET)
        return binary_report

    def __add_current_date_and_time(self):
        date_and_time_formatted = datetime.now().strftime('%d-%m-%Y %H:%M')
        project_name = _('Система моніторингу електроспоживання - KODROS')
        self.__report.add_paragraph(f'{date_and_time_formatted} {project_name}')

    def __add_report_parameters_info(self):
        monitoring_title = self.__report.add_paragraph()
        monitoring_title.add_run(_('\nМоніторинг\n')).bold = True
        query_parameters = self.__query_parameters
        self.__report.add_paragraph(
            _(
                f'Тривалість моніторингу з {query_parameters.period_start} ' + \
                f'по {query_parameters.period_end}'
            )
        )
        monitoring_facility = self.__report.add_paragraph()
        monitoring_facility.add_run(_("Об'єкт моніторингу")).bold = True
        facility = query_parameters.facility_to_get_consumption_for_or_all_descendants_if_any
        if facility.is_institution():
            self.__report.add_paragraph(_(str(facility)))
        else:
            self.__report.add_paragraph(_(str(facility.get_institution())))
            self.__report.add_paragraph(_(str(facility)))
        self.__add_aggregation_interval_info()

    def __add_aggregation_interval_info(self):
        # noinspection PyTypeChecker
        query_parameters = self.__query_parameters
        aggregation_interval_verbose = get_aggregation_interval_verbose(
            query_parameters.aggregation_interval
        )
        self.__report.add_paragraph(_(f'Інтервал агрегації: {aggregation_interval_verbose}'))
        if query_parameters.aggregation_interval == AggregationIntervalSeconds.ONE_HOUR:
            query_parameters: OneHourAggregationIntervalQueryParameters = query_parameters
            if hour_filtering_method := query_parameters.hour_filtering_method:
                hour_filtering_method_verbose = get_hour_filtering_method_verbose(
                    hour_filtering_method
                )
                self.__report.add_paragraph(
                    _(f'Фільтрація інтервалу агрегації: {hour_filtering_method_verbose}')
                )
                self.__report.add_paragraph(
                    _(
                        f'Від: {str(query_parameters.hour_filtering_start_hour).zfill(2)} ' \
                        + f'до: {str(query_parameters.hour_filtering_end_hour).zfill(2)}'
                    )
                )

    def __add_consumption_table(self):
        table_name = self.__report.add_paragraph()
        table_name.add_run(
            _('Відомість фактичних показників споживання електроенергії')
        ).bold = True
        energy_records = [
            record.formatted_consumption_record for record in
            self.__energy_consumption_raw_and_formatted_with_optional_forecast
        ]
        energy_records_count = len(energy_records)
        table = self.__report.add_table(
            cols=self.__CONSUMPTION_TABLE_COLUMNS_COUNT,
            rows=energy_records_count + self.__CONSUMPTION_TABLE_ADDITIONAL_ROWS_COUNT
        )
        table.style = 'Table Grid'
        header_row = table.rows[0]
        INTERVAL_INDEX, VALUE_INDEX = 0, 1
        header_row.cells[INTERVAL_INDEX].text = _('Діапазон')
        header_row.cells[VALUE_INDEX].text = _('Показник електроспоживання, кВт/год')

        for record_number, consumption in enumerate(energy_records, start=1):
            table.rows[record_number].cells[INTERVAL_INDEX].text = consumption.time
            table.rows[record_number].cells[VALUE_INDEX].text = consumption.value

        total_row = table.rows[-1]
        total_text = total_row.cells[0].add_paragraph()
        total_text.add_run(_('Загалом')).bold = True
        total_value = total_row.cells[1].add_paragraph()
        total_value.add_run(self.__total_consumption).bold = True

    def __add_chart(self):
        labels = [
            record.formatted_consumption_record.time for record in
            self.__energy_consumption_raw_and_formatted_with_optional_forecast
        ]
        values = [
            raw_and_formatted_record.raw_consumption_record.value for raw_and_formatted_record in
            self.__energy_consumption_raw_and_formatted_with_optional_forecast
        ]
        __, ax = plt.subplots()
        ax.bar(
            labels, values, color=(1, 0.39, 0.51), hatch='//', label=_('Фактичне споживання'),
            zorder=4, alpha=0.5
        )
        if self.__query_parameters.include_forecast:
            consumption_forecast = [
                record.forecast_raw_and_formatted.raw
                for record in self.__energy_consumption_raw_and_formatted_with_optional_forecast
            ]
            ax.bar(
                labels, consumption_forecast, color=(0.5, 1, 0), hatch=r'\\',
                label=_('Ліміт електроспоживання'), alpha=0.5, zorder=2
            )
        ax.legend()
        plt.grid(zorder=0, axis='y')
        plt.xticks(rotation=90, ha='right')
        plt.tight_layout(pad=2)
        plt.ylabel(_('кВт/год'))
        plt.title(_('Графік споживання теплової енергії'))
        image = io.BytesIO()
        plt.savefig(image, format='png', bbox_inches='tight')
        image.seek(io.SEEK_SET)
        self.__report.add_picture(image)

    def __add_project_manager(self):
        self.__report.add_paragraph(_('Науковий керівник Микола СОТНИК'))
